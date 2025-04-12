"""
Document model for handling document data and operations.
"""
import csv
import logging
import os
import re
from typing import List, Dict, Optional, Set, Tuple, Any

import docx

from models.paragraph import Paragraph, ParaRole

logger = logging.getLogger(__name__)

class Document:
    """Represents a document with paragraphs for Q&A analysis."""
    
    def __init__(self):
        """Initialize an empty document."""
        self.file_path: Optional[str] = None
        self.paragraphs: List[Paragraph] = []
        self.expected_question_count: int = 0
        self._current_q_num: int = 0
        
    def load_file(self, file_path: str, status_callback) -> bool:
        """
        Load a DOCX file, extract paragraphs, run initial analysis.
        
        Args:
            file_path: Path to the DOCX file
            status_callback: Callback function for status updates
            
        Returns:
            bool: True if successful, False otherwise
        """
        try:
            self.file_path = file_path
            status_callback(f"Loading: {os.path.basename(file_path)}...")
            
            # Load document
            doc = docx.Document(self.file_path)
            raw_paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and not p.text.isspace()]
            logger.info(f"Extracted {len(raw_paragraphs)} non-empty paragraphs.")
            
            if not raw_paragraphs:
                logger.error("Document contains no readable text.")
                return False
                
            # Run initial analysis
            return self._analyze_document(raw_paragraphs, status_callback)
            
        except Exception as e:
            logger.error(f"Error loading or processing file: {e}", exc_info=True)
            return False
    
    def _analyze_document(self, raw_paragraphs: List[str], status_callback) -> bool:
        """
        Analyze document content and identify questions and answers.
        
        Args:
            raw_paragraphs: List of paragraph texts
            status_callback: Callback function for status updates
            
        Returns:
            bool: True if successful, False otherwise
        """
        try:
            # Get initial questions and estimate count
            question_indices, est_count = self._get_initial_question_indices(raw_paragraphs, status_callback)
            self.expected_question_count = est_count
            
            # Process paragraphs
            self._process_paragraphs(raw_paragraphs, question_indices)
            
            return True
        except Exception as e:
            logger.error(f"Error analyzing document: {e}", exc_info=True)
            return False
    
    def _process_paragraphs(self, raw_paragraphs: List[str], question_indices: Set[int]) -> None:
        """
        Process paragraphs and assign roles based on initial analysis.
        
        Args:
            raw_paragraphs: List of paragraph texts
            question_indices: Set of indices identified as questions
        """
        self.paragraphs = []
        self._current_q_num = 0
        last_role = ParaRole.UNDETERMINED
        
        for i, text in enumerate(raw_paragraphs):
            role = ParaRole.UNDETERMINED
            q_num = None
            
            if i in question_indices:
                role = ParaRole.QUESTION
                self._current_q_num += 1
                q_num = self._current_q_num
            elif last_role == ParaRole.QUESTION or last_role == ParaRole.ANSWER:
                # If the previous was Q or A, assume this is an Answer
                role = ParaRole.ANSWER
                q_num = self._current_q_num
            else:
                # Could be header or undetermined
                # Mark short starting lines as IGNORE potentially
                if i < 5 and len(text) < 50:  # Crude header check
                    role = ParaRole.IGNORE
                else:
                    role = ParaRole.UNDETERMINED
            
            paragraph = Paragraph(index=i, text=text, role=role, q_num=q_num)
            self.paragraphs.append(paragraph)
            last_role = role
    
    def _get_initial_question_indices(self, paragraphs: List[str], status_callback) -> Tuple[Set[int], int]:
        """
        Uses heuristic scoring to identify potential question paragraphs.
        
        Args:
            paragraphs: List of paragraph texts
            status_callback: Callback function for status updates
            
        Returns:
            Tuple containing set of question indices and estimated question count
        """
        status_callback("Running initial heuristic analysis...")
        
        # First phase: Identify patterns to estimate question count
        estimated_count = self._estimate_question_count(paragraphs, status_callback)
        status_callback(f"Estimated question count from document structure: {estimated_count}")
        
        potential_questions = []
        # Map original index to paragraph text for scoring
        para_map = {i: p for i, p in enumerate(paragraphs)}

        # Identify potential questions based on keywords/structure
        for i, p in enumerate(paragraphs):
            if len(p) < 10: continue
            if p.startswith(('```', '<!--', '-->')): continue
            
            has_question_mark = '?' in p
            has_question_words = any(word in p.lower() for word in
                             ['what', 'when', 'where', 'why', 'how', 'name', 'list',
                              'identify', 'describe', 'define', 'explain'])
            has_numeric_reference = bool(re.search(r'(\d+)\s+(kinds|types|requirements|grounds|factors|situations|matters|things)', p, re.IGNORECASE))
            is_question_like = (
                p.lower().startswith(('what', 'when', 'where', 'why', 'how', 'name', 'is ', 'are ', 'does ')) or
                re.search(r'(what is|what are|name the|which|how many)', p.lower())
            )

            if has_question_mark or has_question_words or has_numeric_reference or is_question_like:
                potential_questions.append((i, p))

        status_callback(f"Found {len(potential_questions)} potential questions based on initial keywords/structure.")

        # Score the potential questions
        scored_questions = []
        for idx, text in potential_questions:
            score = 0
            if '?' in text: score += 10
            if re.match(r'^(What|When|Where|Why|How|Name|Which|Is|Are|Does|Do)\b', text, re.IGNORECASE): score += 8
            if any(word in text.lower() for word in ['what', 'when', 'where', 'why', 'how']): score += 5
            if any(word in text.lower() for word in ['name', 'list', 'identify', 'describe', 'define']): score += 5
            if re.search(r'(\d+)\s+(kinds|types|requirements|grounds|factors|situations|matters|things)', text, re.IGNORECASE): score += 7
            if len(text) > 30: score += 3
            # Simplified next paragraph check - less reliable, weighted lower
            if idx < len(paragraphs) - 1 and re.match(r'^\s*\d+[\.\)]', para_map.get(idx + 1, '')): score += 2
            if text.startswith(('The ', 'A ', 'An ', 'Both ', 'It ', 'PC', 'BRO', 'LAC', '1.', '2.', '3.', 'a.', 'b.')): score -= 5 # Penalize typical answer starts
            if text.isupper() or text.startswith('**'): score -= 3

            # Big boost if it *actually* starts with a number/dot/space pattern
            if re.match(r"^\s*\d+\s*[\.\)]\s+", text):
                score += 15

            # Only add if score is reasonably positive
            if score > 0:
               scored_questions.append({'index': idx, 'text': text, 'score': score})

        scored_questions.sort(key=lambda x: x['score'], reverse=True)
        status_callback(f"Scored {len(scored_questions)} potential questions.")

        # Take the top candidates based on the estimated count, but be slightly generous
        take_count = min(len(scored_questions), max(estimated_count, int(estimated_count * 1.1)))
        top_scored_indices = {q['index'] for q in scored_questions[:take_count]}

        status_callback(f"Identified initial {len(top_scored_indices)} candidates for questions.")
        return top_scored_indices, estimated_count
    
    def _estimate_question_count(self, paragraphs: List[str], status_callback) -> int:
        """
        Estimates the number of questions in the document based on structure.
        
        Args:
            paragraphs: List of paragraph texts
            status_callback: Callback function for status updates
            
        Returns:
            int: Estimated question count
        """
        # Method 1: Look for sequential numbering patterns
        numbered_paragraphs = [p for p in paragraphs if re.match(r'^\s*\d+[\.\)]', p)]
        
        # Method 2: Look for question marks
        question_marks = [p for p in paragraphs if '?' in p]
        
        # Method 3: Check document title/header for question count
        title_count = None
        for i, p in enumerate(paragraphs[:10]):  # Check first 10 paragraphs for headers
            count_match = re.search(r'(\d+)\s*(?:questions|problems|items)', p, re.IGNORECASE)
            if count_match:
                title_count = int(count_match.group(1))
                status_callback(f"Found question count in document header: {title_count}")
                break
        
        # Method 4: Look at largest sequential number
        max_seq_num = 0
        for p in paragraphs:
            num_match = re.match(r'^\s*(\d+)[\.\)]', p)
            if num_match:
                try:
                    num = int(num_match.group(1))
                    max_seq_num = max(max_seq_num, num)
                except ValueError:
                    pass
        
        # Decide on the most likely count
        if title_count is not None:
            return title_count
        elif max_seq_num > 10:  # If we have sequential numbering with a reasonable max
            return max_seq_num
        elif len(numbered_paragraphs) > 10:
            return len(numbered_paragraphs)
        elif len(question_marks) > 5:
            return len(question_marks)
        else:
            # Default fallback - use a reasonable default
            return 25
    
    def renumber_questions(self) -> None:
        """Renumber questions and answers sequentially."""
        q_counter = 0
        current_q_num_for_answers = 0
        
        for para in self.paragraphs:
            if para.role == ParaRole.QUESTION:
                q_counter += 1
                para.q_num = q_counter
                current_q_num_for_answers = q_counter
            elif para.role == ParaRole.ANSWER:
                # Assign answer to the most recently seen question number
                para.q_num = current_q_num_for_answers
            else:  # IGNORE or UNDETERMINED
                para.q_num = None
        
        self._current_q_num = q_counter
    
    def change_paragraph_role(self, index: int, new_role: ParaRole) -> bool:
        """
        Change the role of a paragraph.
        
        Args:
            index: Index of the paragraph to change
            new_role: New role to assign
            
        Returns:
            bool: True if renumbering is needed, False otherwise
        """
        if 0 <= index < len(self.paragraphs):
            old_role = self.paragraphs[index].role
            self.paragraphs[index].role = new_role
            
            # If changing to/from QUESTION, we need to renumber
            if old_role == ParaRole.QUESTION or new_role == ParaRole.QUESTION:
                return True
        return False
    
    def merge_paragraph_up(self, index: int) -> bool:
        """
        Merge a paragraph into the previous answer block.
        
        Args:
            index: Index of the paragraph to merge
            
        Returns:
            bool: True if renumbering is needed, False otherwise
        """
        if index <= 0 or index >= len(self.paragraphs):
            return False
            
        # Find the effective q_num of the *preceding* block
        preceding_q_num = None
        for prev_idx in range(index - 1, -1, -1):
            if self.paragraphs[prev_idx].q_num is not None:
                preceding_q_num = self.paragraphs[prev_idx].q_num
                break

        if preceding_q_num is not None:
            old_role = self.paragraphs[index].role
            self.paragraphs[index].role = ParaRole.ANSWER
            self.paragraphs[index].q_num = preceding_q_num
            
            # If it was a QUESTION, we need renumbering
            if old_role == ParaRole.QUESTION:
                return True
        return False
    
    def get_qa_data(self) -> Tuple[Dict[int, Dict[str, Any]], int]:
        """
        Get structured Q&A data for export.
        
        Returns:
            Tuple containing dictionary of QA pairs and question count
        """
        questions_data = {}  # {q_num: {'question': text, 'answers': [text]}}
        q_count = 0

        # First pass: Collect questions
        for para in self.paragraphs:
            if para.role == ParaRole.QUESTION:
                q_num = para.q_num
                if q_num is not None:
                    q_count += 1
                    questions_data[q_num] = {'number': q_num, 'question': para.text, 'answers': []}

        # Second pass: Collect answers
        for para in self.paragraphs:
            if para.role == ParaRole.ANSWER and para.q_num in questions_data:
                questions_data[para.q_num]['answers'].append(para.text)

        return questions_data, q_count
    
    def save_to_csv(self, save_path: str) -> bool:
        """
        Save the verified Q&A pairs to a CSV file.
        
        Args:
            save_path: Path to save the CSV file
            
        Returns:
            bool: True if successful, False otherwise
        """
        try:
            questions_data, _ = self.get_qa_data()
            
            with open(save_path, 'w', newline='', encoding='utf-8') as f:
                writer = csv.writer(f)
                for q_num in sorted(questions_data.keys()):
                    q_data = questions_data[q_num]
                    # Format: Question Number. Question Text in first column
                    # Answer paragraphs in subsequent columns
                    row = [f"{q_data['number']}. {q_data['question']}"] + q_data['answers']
                    writer.writerow(row)
            
            logger.info(f"Successfully saved verified data to: {save_path}")
            return True
            
        except Exception as e:
            logger.error(f"Failed to write CSV file: {e}", exc_info=True)
            return False

    def get_question_count(self) -> int:
        """Get the current number of questions in the document."""
        return sum(1 for p in self.paragraphs if p.role == ParaRole.QUESTION)
    
    def set_expected_question_count(self, count: int) -> None:
        """Set the expected question count."""
        if count > 0:
            self.expected_question_count = count