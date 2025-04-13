"""
File service for handling file operations.
"""
import csv
import logging
import os
import platform
import subprocess
import threading
from tkinter import filedialog
from typing import List, Optional, Tuple, Dict, Any, Callable

import docx

logger = logging.getLogger(__name__)

class FileService:
    """Service for handling file operations."""
    
    @staticmethod
    def select_docx_file() -> Optional[str]:
        """
        Open file dialog to select a DOCX file.
        
        Returns:
            str: Selected file path or None if cancelled
        """
        path = filedialog.askopenfilename(
            title="Select DOCX File",
            filetypes=[("Word Documents", "*.docx")]
        )
        
        if path:
            logger.info(f"Selected file: {path}")
            return path
        
        logger.info("File selection cancelled.")
        return None
    
    @staticmethod
    def get_save_csv_path(original_file_path: str) -> Optional[str]:
        """
        Get save path for CSV export.
        
        Args:
            original_file_path: Original file path
            
        Returns:
            str: Save path or None if cancelled
        """
        if not original_file_path:
            return None
            
        # Generate default filename
        output_csv_path = os.path.splitext(original_file_path)[0] + "_verified.csv"
        
        save_path = filedialog.asksaveasfilename(
            title="Save Verified CSV",
            initialfile=os.path.basename(output_csv_path),
            defaultextension=".csv",
            filetypes=[("CSV files", "*.csv")]
        )
        
        if save_path:
            logger.info(f"Selected save path: {save_path}")
            return save_path
            
        logger.info("CSV save cancelled by user.")
        return None
    
    @staticmethod
    def open_file_externally(file_path: str) -> Tuple[bool, str]:
        """
        Open a file using the system's default application.
        
        Args:
            file_path: Path to the file to open
            
        Returns:
            Tuple containing success flag and error message if any
        """
        try:
            if platform.system() == 'Windows':
                os.startfile(file_path)
            elif platform.system() == 'Darwin':  # macOS
                subprocess.call(['open', file_path])
            else:  # Linux and others
                subprocess.call(['xdg-open', file_path])
                
            logger.info(f"Opened file externally: {file_path}")
            return True, ""
        except Exception as e:
            error_msg = f"Could not open file: {str(e)}"
            logger.error(error_msg)
            return False, error_msg
            
    @staticmethod
    def load_docx_paragraphs(file_path: str) -> List[str]:
        """
        Load paragraphs from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            List of paragraph texts
        """
        try:
            # Load document
            doc = docx.Document(file_path)
            raw_paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and not p.text.isspace()]
            logger.info(f"Extracted {len(raw_paragraphs)} non-empty paragraphs from: {file_path}")
            
            return raw_paragraphs
            
        except Exception as e:
            logger.error(f"Error loading DOCX file: {e}", exc_info=True)
            raise
    
    @staticmethod
    def load_docx_paragraphs_async(file_path: str, callback: Callable[[List[str], Optional[Exception]], None]) -> threading.Thread:
        """
        Load paragraphs from a DOCX file asynchronously.
        
        Args:
            file_path: Path to the DOCX file
            callback: Callback function receiving (paragraphs, exception) upon completion
            
        Returns:
            Thread object
        """
        def _load_thread():
            try:
                # Load document
                doc = docx.Document(file_path)
                raw_paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and not p.text.isspace()]
                logger.info(f"Extracted {len(raw_paragraphs)} non-empty paragraphs from: {file_path}")
                
                # Call callback with results and no exception
                callback(raw_paragraphs, None)
                
            except Exception as e:
                logger.error(f"Error loading DOCX file: {e}", exc_info=True)
                # Call callback with no results and the exception
                callback(None, e)
        
        # Create and start thread
        thread = threading.Thread(target=_load_thread)
        thread.daemon = True
        thread.start()
        
        return thread
    
    @staticmethod
    def save_data_to_csv(data: List[List[str]], save_path: str) -> bool:
        """
        Save data to a CSV file.
        
        Args:
            data: List of data rows
            save_path: Path to save the CSV file
            
        Returns:
            bool: True if successful, False otherwise
        """
        try:
            with open(save_path, 'w', newline='', encoding='utf-8') as f:
                writer = csv.writer(f)
                for row in data:
                    writer.writerow(row)
            
            logger.info(f"Successfully saved data to: {save_path}")
            return True
            
        except Exception as e:
            logger.error(f"Failed to write CSV file: {e}", exc_info=True)
            return False